# ==========================================
# MUST BE FIRST (Fix oneDNN / MKL crash)
# ==========================================
import html
import os
import re

os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["KMP_INIT_AT_FORK"] = "FALSE"

# ==========================================
# Imports
# ==========================================

import time
import pymupdf as fitz
import numpy as np
import cv2
import gc
import ctypes
from multiprocessing import Process, Queue
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from pathlib import Path
from typing import Optional

_ws = re.compile(r"\s+")


def normalize_ocr_text(text: str) -> str:
    return _ws.sub(" ", text.replace("\u00A0", " ")).strip()


def create_paddle_ocr() -> PaddleOCR:
    return PaddleOCR(
        use_angle_cls=False,
        lang="en",
        use_gpu=False,
        enable_mkldnn=False,
        cpu_threads=2,
        rec_batch_num=2,
    )


def normalize_legal_text(text: str) -> str:
    # normalize whitespace
    text = text.replace("\u00A0", " ")
    text = _ws.sub(" ", text).strip()

    fixes = [
        (r"\bU\s*\.\s*S\s*\.\s*C\s*\.\b", "U.S.C."),
        (r"\bC\s*\.\s*F\s*\.\s*R\s*\.\b", "C.F.R."),
        (r"\bI\s*&\s*N\b", "I&N"),
        (r"\bI\s*&\s*N\s+Dec\.\b", "I&N Dec."),
        (r"\bI\s*&\s*N\s+Dec\b", "I&N Dec."),
        (r"\b1NA\b", "INA"),
        (r"\b8\s*C\s*\.\s*F\s*\.\s*R\s*\.\b", "8 C.F.R."),
        (r"\b8\s+U\s*\.\s*S\s*\.\s*C\s*\.\b", "8 U.S.C."),
        (r"\(\s*IJ\s+at\s+(\d+)\s*\)", r"(IJ at \1)"),
    ]
    for pat, repl in fixes:
        text = re.sub(pat, repl, text)

    # restore section symbol when OCR emits 'S' or '$' after citations
    text = re.sub(r"\b(U\.S\.C\.|C\.F\.R\.|INA)\s+[Ss\$]\s+(\d[\w.\-]*)", r"\1 § \2", text)
    text = re.sub(r"\b(U\.S\.C\.|C\.F\.R\.)\s+[Ss\$]{2}\s+(\d[\w.\-]*)", r"\1 §§ \2", text)
    text = re.sub(r"\b(Section)\s+[Ss\$]\s+(\d[\w.\-]*)", r"\1 § \2", text, flags=re.IGNORECASE)

    # plain s / $ used as section symbol
    text = re.sub(r"\b[sS]\s+(\d[\w.\-]*)", r"§ \1", text)
    text = re.sub(r"\$\s+(\d[\w.\-]*)", r"§ \1", text)
    # double section symbols
    text = re.sub(r"\b[sS]{2}\s+(\d[\w.\-]*)", r"§§ \1", text)
    text = re.sub(r"\$\$\s+(\d[\w.\-]*)", r"§§ \1", text)

    # normalize spacing around section symbol
    text = re.sub(r"\s*§\s*", " § ", text)
    text = re.sub(r"§\s+§", "§§", text)
    text = _ws.sub(" ", text).strip()

    # plain s / $ / g / j used as section symbol
    text = re.sub(r"\b[sSgGjJ]\s+(\d[\w.\-]*)", r"§ \1", text)
    text = re.sub(r"\$\s+(\d[\w.\-]*)", r"§ \1", text)

    # double section symbols (ss / gg / gs / $$)
    text = re.sub(r"\b(?:[sS]{2}|[gG]{2}|[gG][sS])\s+(\d[\w.\-]*)", r"§§ \1", text)
    text = re.sub(r"\$\$\s+(\d[\w.\-]*)", r"§§ \1", text)
    text = re.sub(r"\b(U\.S\.C\.|C\.F\.R\.)\s*(\d{1,4})\b", r"\1 § \2", text)

    # normalize spacing around section symbol
    text = re.sub(r"\s*§\s*", " § ", text)
    text = re.sub(r"§\s+§", r"§§", text)
    text = _ws.sub(" ", text).strip()

    return text


# ==========================================
# Image Preprocessing (Balanced Quality/Speed)
# ==========================================
def preprocess_image(img_rgb):
    """
    Improve OCR for thin characters like 'ii' vs 'i':
    - upscale a bit
    - convert to gray
    - mild denoise
    - contrast boost (CLAHE)
    - adaptive threshold
    """
    # upscale (helps distinguish ii)
    h, w = img_rgb.shape[:2]
    scale = 1.5
    img_rgb = cv2.resize(img_rgb, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)

    gray = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)

    # denoise but keep edges
    gray = cv2.bilateralFilter(gray, 7, 50, 50)

    # contrast boost
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # binarize
    thresh = cv2.adaptiveThreshold(
        gray, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31, 9
    )
    return thresh


def extract_text_from_image(
    image_path: str | Path,
    ocr: Optional[PaddleOCR] = None,
    *,
    normalize_legal: bool = False,
) -> str:
    owns_ocr = ocr is None
    ocr_instance = ocr or create_paddle_ocr()
    image = cv2.imread(str(image_path))
    if image is None:
        raise FileNotFoundError(f"Could not read image: {image_path}")

    rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    prepared = preprocess_image(rgb_image)
    result = ocr_instance.ocr(prepared, cls=False)
    if not result or not result[0]:
        if owns_ocr:
            del ocr_instance
        return ""

    text_lines: list[str] = []
    for line in order_paddleocr_boxes_reading_order(result[0]):
        parts = []
        for _, _, _, _, box in line:
            candidate = normalize_ocr_text(box[1][0] or "")
            if normalize_legal and candidate:
                candidate = normalize_legal_text(candidate)
            if candidate:
                parts.append(candidate)
        if parts:
            text_lines.append(" ".join(parts))

    if owns_ocr:
        del ocr_instance

    return "\n".join(text_lines)


def extract_text_from_images(
    image_files,
    ocr: Optional[PaddleOCR] = None,
    *,
    escape_html: bool = False,
    normalize_legal: bool = False,
) -> str:
    owns_ocr = ocr is None
    ocr_instance = ocr or create_paddle_ocr()
    text = "\n".join(
        extracted
        for extracted in (
            extract_text_from_image(image_file, ocr_instance, normalize_legal=normalize_legal)
            for image_file in image_files
        )
        if extracted
    )

    if owns_ocr:
        del ocr_instance

    if escape_html:
        return html.escape(text, quote=False)
    return text


def _write_text_file(text: str, text_path: str | Path) -> None:
    output_path = Path(text_path)
    output_path.write_text(text.rstrip() + "\n" if text else "", encoding="utf-8")


# ==========================================
# Optional .env loader (no extra deps)
# ==========================================
def load_dotenv_if_present(env_path: Optional[Path] = None) -> None:
    """Load key=value pairs from a local .env file into os.environ.

    Existing environment variables are NOT overwritten.
    """
    if env_path is None:
        env_path = Path(__file__).resolve().parent / ".env"
    if not env_path.exists():
        return
    try:
        for raw in env_path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            k = k.strip()
            v = v.strip().strip('"').strip("'")
            if k and k not in os.environ:
                os.environ[k] = v
    except Exception:
        return


# ==========================================
# PDF text-layer extraction (preserve italics)
# ==========================================
def is_span_italic(span: dict) -> bool:
    font_name = (span.get("font") or "").lower()
    if "italic" in font_name or "oblique" in font_name:
        return True
    flags = span.get("flags")
    try:
        flags = int(flags)
    except Exception:
        flags = 0
    return bool(flags & 2)


def pdf_page_to_docx_runs(page: "fitz.Page", doc: Document) -> bool:
    """Write page text into docx preserving italics; returns True if any text found."""
    d = page.get_text("dict")
    blocks = d.get("blocks", [])
    wrote_any = False

    for b in blocks:
        if b.get("type") != 0:
            continue
        for line in b.get("lines", []):
            spans = line.get("spans", [])
            para = doc.add_paragraph()
            has_text = False
            for span in spans:
                text = span.get("text") or ""
                if not text.strip():
                    continue
                run = para.add_run(text)
                run.italic = is_span_italic(span)
                has_text = True
            if has_text:
                wrote_any = True

    return wrote_any


# ==========================================
# Memory Cleanup
# ==========================================
def force_memory_cleanup():
    gc.collect()
    try:
        ctypes.windll.kernel32.SetProcessWorkingSetSize(
            ctypes.windll.kernel32.GetCurrentProcess(),
            -1,
            -1
        )
    except:
        pass


def _flush_process_memory() -> None:
    force_memory_cleanup()


# ==========================================
# Worker: Process One PDF
# ==========================================
def _box_metrics(box):
    pts = box[0]  # 4 points
    xs = [p[0] for p in pts]
    ys = [p[1] for p in pts]
    left, right = min(xs), max(xs)
    top, bottom = min(ys), max(ys)
    h = max(1.0, bottom - top)
    yc = (top + bottom) / 2.0
    return left, right, top, bottom, h, yc


def order_paddleocr_boxes_reading_order(raw_boxes):
    """
    Convert PaddleOCR boxes into reading-order lines:
    1) cluster by y-center into lines
    2) sort each line by left x
    Returns: list of lines, where each line is list of box objects.
    """
    items = []
    for b in raw_boxes:
        text = (b[1][0] or "").strip()
        if not text:
            continue
        left, right, top, bottom, h, yc = _box_metrics(b)
        items.append((yc, top, bottom, left, h, b))

    if not items:
        return []

    # Sort by y-center then x-left
    items.sort(key=lambda t: (t[0], t[3]))

    # Robust line threshold based on typical box height
    heights = sorted(t[4] for t in items)
    median_h = heights[len(heights) // 2] if heights else 12.0
    line_y_thresh = max(8.0, 0.55 * median_h)

    lines = []
    current = []
    current_y = None

    for yc, top, bottom, left, h, b in items:
        if current_y is None:
            current = [(yc, top, bottom, left, b)]
            current_y = yc
            continue

        if abs(yc - current_y) <= line_y_thresh:
            current.append((yc, top, bottom, left, b))
            # slowly adapt the running y-center
            current_y = (current_y * 0.8) + (yc * 0.2)
        else:
            # finalize current line
            current.sort(key=lambda t: t[3])  # x-left
            lines.append(current)
            # start new line
            current = [(yc, top, bottom, left, b)]
            current_y = yc

    if current:
        current.sort(key=lambda t: t[3])
        lines.append(current)

    return lines


def process_pdf_worker(pdf_path, docx_path, queue=None):
    try:
        pdf = fitz.open(pdf_path)
    except Exception as e:
        print(f"Could not open {pdf_path}: {e}")
        if queue:
            queue.put((pdf_path, "fail"))
        return

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(14)

    # Prefer extracting PDF text layer first to preserve italics.
    prefer_text_layer = os.environ.get("PADDLE_PREFER_PDF_TEXT_LAYER", "1").strip() not in {"0", "false", "False"}
    ocr = None

    for page_number in range(len(pdf)):
        try:
            page = pdf[page_number]

            # Add page break BEFORE next page (match existing behavior)
            if page_number > 0:
                doc.add_page_break()

            if prefer_text_layer:
                wrote_any = pdf_page_to_docx_runs(page, doc)
                if wrote_any:
                    # No OCR needed for this page
                    del page
                    continue

            # Fall back to OCR for this page
            if ocr is None:
                ocr = create_paddle_ocr()

            pix = page.get_pixmap(dpi=400)
            print("*" * 60)
            print("DPI: 400")
            print("*" * 60)

            img = np.frombuffer(
                pix.samples, dtype=np.uint8
            ).reshape(pix.height, pix.width, pix.n)

            if pix.n == 4:
                img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)

            img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

            img_for_ocr = preprocess_image(img)
            result = ocr.ocr(img_for_ocr, cls=False)

            if result and result[0]:

                lines = order_paddleocr_boxes_reading_order(result[0])

                paragraph_text = ""
                prev_line_bottom = None
                prev_line_top = None

                # threshold for paragraph breaks
                # (based on median height-ish; we approximate using first box height per line)
                for line in lines:
                    # line bounds
                    line_tops = []
                    line_bottoms = []
                    line_text_parts = []

                    for (_, top, bottom, _, box) in line:
                        text = (box[1][0] or "").strip()
                        if not text:
                            continue

                        text = re.sub(
                            r"(?:\(\s*)?(?:A\s*\)?\s*)?\(?\s*b\s*\)?\s*\(?\s*6\s*\)?",
                            "[redacted]",
                            text,
                            flags=re.IGNORECASE
                        )

                        text = re.sub(
                            r"""
                            \( \s* b \s* \\ \s* \( \s* 6 \s* \)? \s*
                            |
                            \( \s* A \s* \( \s* h \s* r \s* 6 \s* \) \s* \) \s*
                            |
                            \( \s* h \s* \\ \s* \( \s* 6 \s* \)? \s*
                            """,
                            "[redacted]",
                            text,
                            flags=re.IGNORECASE | re.VERBOSE
                        )

                        text = normalize_legal_text(text)

                        line_tops.append(top)
                        line_bottoms.append(bottom)
                        line_text_parts.append(text)

                    if not line_text_parts:
                        continue

                    line_top = min(line_tops)
                    line_bottom = max(line_bottoms)
                    line_text = " ".join(line_text_parts)

                    # Decide paragraph break by vertical gap from previous line
                    if prev_line_bottom is not None:
                        gap = line_top - prev_line_bottom

                        # dynamic-ish threshold: big gaps create new paragraphs
                        if gap >= 24:
                            if paragraph_text.strip():
                                doc.add_paragraph(paragraph_text.strip())
                            paragraph_text = line_text
                        else:
                            # same paragraph
                            paragraph_text = (paragraph_text + " " + line_text).strip()
                    else:
                        paragraph_text = line_text

                    prev_line_bottom = line_bottom
                    prev_line_top = line_top

                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            del img, img_for_ocr, result, pix, page
            _flush_process_memory()

        except Exception as e:
            print(f"Error on page {page_number + 1}: {e}")
            continue

    pdf.close()

    try:
        doc.save(docx_path)
        print(f"Saved: {docx_path}")
        if queue:
            queue.put((pdf_path, "success"))
    except Exception as e:
        print(f"Failed saving {docx_path}: {e}")
        if queue:
            queue.put((pdf_path, "fail"))

    del doc
    if ocr is not None:
        del ocr
    _flush_process_memory()


def process_pdf_file(
    input_file: str | Path,
    output_folder: str | Path,
    queue: Optional[Queue] = None,
) -> Path:
    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"PDF file does not exist: {input_path}")
    if not input_path.is_file():
        raise FileNotFoundError(f"Input path is not a file: {input_path}")
    if input_path.suffix.lower() != ".pdf":
        raise ValueError(f"Input file must be a PDF: {input_path}")

    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    docx_path = output_path / f"{input_path.stem}.docx"
    process_pdf_worker(str(input_path), str(docx_path), queue)
    return docx_path


def _sanitize_output_part(part: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]+', "_", part.strip())
    sanitized = re.sub(r"\s+", " ", sanitized).strip()
    return sanitized


def _get_pdf_output_stem(pdf_path: str | Path, input_root: str | Path) -> str:
    input_path = Path(pdf_path)
    root_path = Path(input_root)

    relative_path = input_path.relative_to(root_path).with_suffix("")
    parts = [_sanitize_output_part(part) for part in relative_path.parts if _sanitize_output_part(part)]
    return "_".join(parts)


def _get_pdf_text_output_path(
    pdf_path: str | Path,
    output_folder: str | Path,
    input_root: str | Path,
) -> Path:
    return Path(output_folder) / f"{_get_pdf_output_stem(pdf_path, input_root)}.txt"


def extract_text_from_pdf(pdf_path: str | Path) -> str:
    pdf = fitz.open(str(pdf_path))
    prefer_text_layer = os.environ.get("PADDLE_PREFER_PDF_TEXT_LAYER", "1").strip() not in {"0", "false", "False"}
    ocr = None
    page_texts: list[str] = []

    try:
        for page in pdf:
            page_text = ""

            if prefer_text_layer:
                page_text = (page.get_text("text") or "").strip()

            if not page_text:
                if ocr is None:
                    ocr = create_paddle_ocr()

                pix = page.get_pixmap(dpi=400)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

                if pix.n == 4:
                    img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)

                img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                img_for_ocr = preprocess_image(img)
                result = ocr.ocr(img_for_ocr, cls=False)

                if result and result[0]:
                    line_texts = []
                    for line in order_paddleocr_boxes_reading_order(result[0]):
                        parts = []
                        for _, _, _, _, box in line:
                            candidate = normalize_ocr_text(box[1][0] or "")
                            if candidate:
                                parts.append(normalize_legal_text(candidate))
                        if parts:
                            line_texts.append(" ".join(parts))
                    page_text = "\n".join(line_texts).strip()

                del img, img_for_ocr, result, pix
                _flush_process_memory()

            if page_text:
                page_texts.append(page_text)
    finally:
        pdf.close()
        if ocr is not None:
            del ocr
        _flush_process_memory()

    return fix_common_cfr_subsections("\n\n".join(page_texts).strip())


def process_pdf_to_text_file(
    input_file: str | Path,
    output_folder: str | Path,
    input_root: str | Path,
) -> Path:
    input_path = Path(input_file)
    if not input_path.exists():
        raise FileNotFoundError(f"PDF file does not exist: {input_path}")
    if not input_path.is_file():
        raise FileNotFoundError(f"Input path is not a file: {input_path}")
    if input_path.suffix.lower() != ".pdf":
        raise ValueError(f"Input file must be a PDF: {input_path}")

    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    text_output_path = _get_pdf_text_output_path(input_path, output_path, input_root)
    extracted_text = extract_text_from_pdf(input_path)
    _write_text_file(extracted_text, text_output_path)
    return text_output_path


def _process_pdf_to_text_file_worker(
    input_file: str | Path,
    output_folder: str | Path,
    input_root: str | Path,
    queue: Queue,
) -> None:
    output_file: Optional[Path] = None
    try:
        output_file = process_pdf_to_text_file(input_file, output_folder, input_root)
        queue.put(("success", str(output_file), None))
    except Exception as exc:
        queue.put(("fail", str(input_file), str(exc)))
    finally:
        if output_file is not None:
            del output_file
        _flush_process_memory()


def process_pdf_to_text_file_isolated(
    input_file: str | Path,
    output_folder: str | Path,
    input_root: str | Path,
    *,
    timeout_seconds: int = 600,
) -> Path:
    queue = Queue()
    process = Process(
        target=_process_pdf_to_text_file_worker,
        args=(str(input_file), str(output_folder), str(input_root), queue),
    )
    process.start()
    process.join(timeout=timeout_seconds)

    if process.is_alive():
        process.terminate()
        process.join()
        raise TimeoutError(f"Timed out converting PDF to text: {input_file}")

    if queue.empty():
        raise RuntimeError(f"No result returned for PDF conversion: {input_file}")

    status, output_value, error_message = queue.get()
    queue.close()
    queue.join_thread()

    if status != "success":
        raise RuntimeError(error_message or f"Failed converting PDF to text: {input_file}")

    _flush_process_memory()
    return Path(output_value)


def pdf_to_text_ocr_main(input_folder: str | Path, output_folder: str | Path) -> list[Path]:
    print("PDF OCR Started\n")

    load_dotenv_if_present()

    input_path = Path(input_folder)
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        raise FileNotFoundError(f"Input folder does not exist: {input_path}")
    if not input_path.is_dir():
        raise NotADirectoryError(f"Input path is not a folder: {input_path}")

    pdf_files = sorted(path for path in input_path.rglob("*.pdf") if path.is_file())
    if not pdf_files:
        raise FileNotFoundError(f"No PDF files found in: {input_path}")

    output_files: list[Path] = []
    successes = 0
    fails = 0
    skips = 0
    start_time = time.time()
    timeout_seconds = int(os.environ.get("PADDLE_PDF_TIMEOUT_SECONDS", "600"))

    for pdf_file in pdf_files:
        try:
            print(f"Processing PDF: {pdf_file.relative_to(input_path)}")
            output_file = _get_pdf_text_output_path(pdf_file, output_path, input_path)
            if output_file.exists():
                skips += 1
                output_files.append(output_file)
                print(f"Skipping existing text file: {output_file}")
                continue

            output_file = process_pdf_to_text_file_isolated(
                pdf_file,
                output_path,
                input_path,
                timeout_seconds=timeout_seconds,
            )
            output_files.append(output_file)
            successes += 1
            print(f"Saved text: {output_file}")
        except Exception as exc:
            fails += 1
            print(f"Failed PDF {pdf_file}: {exc}")
        finally:
            _flush_process_memory()

    total_seconds = time.time() - start_time

    print("\n==============================")
    print("DONE")
    print("==============================")
    print(f"Success: {successes}")
    print(f"Skipped: {skips}")
    print(f"Fail: {fails}")
    print("Output files:")
    for output_file in output_files:
        print(output_file)
    print(f"Total time: {total_seconds / 60:.2f} minutes")

    return output_files


def fix_common_cfr_subsections(text: str) -> str:
    """
    Fix OCR confusion i vs ii ONLY in the exact context of:
    8 C.F.R. § 1003.1(d)(3)(i/ii)
    """
    t = text

    t = re.sub(
        r"(8\s*C\.F\.R\.\s*§\s*1003\.1\s*\(\s*d\s*\)\s*\(\s*3\s*\)\s*)\(\s*i\s*\)",
        r"\1(ii)",
        t,
        flags=re.IGNORECASE
    )
    return t


# ==========================================
# Main Pipeline
# ==========================================
def main():
    print("OCR Batch Started\n")

    load_dotenv_if_present()

    input_folder = os.environ.get("PADDLE_INPUT_FOLDER", r"C:\Users\tior\Downloads\OneDrive_1_4-24-2026\Rel 133")
    output_folder = os.environ.get("PADDLE_OUTPUT_FOLDER", r"C:\Users\tior\Downloads\OneDrive_1_4-24-2026\Rel 133\output")

    os.makedirs(output_folder, exist_ok=True)

    start_time = time.time()

    pdf_files = [f for f in os.listdir(input_folder) if f.lower().endswith(".pdf")]

    if not pdf_files:
        print("No PDFs found.")
        return

    print(f"Found {len(pdf_files)} PDFs\n")

    queue = Queue()
    processes = []

    for file_name in pdf_files:
        pdf_path = os.path.join(input_folder, file_name)

        print(f"Starting process for {file_name}")

        p = Process(target=process_pdf_file, args=(pdf_path, output_folder, queue))
        p.start()
        processes.append(p)

        max_processes = int(os.environ.get("PADDLE_MAX_PROCESSES", "2"))
        if len(processes) >= max_processes:
            for proc in processes:
                print("Waiting for process:", proc.pid)
                proc.join(timeout=600)
                if proc.is_alive():
                    print("Process stuck, terminating:", proc.pid)
                    proc.terminate()
                    proc.join()
            processes = []

    for proc in processes:
        print("Waiting for process:", proc.pid)
        proc.join(timeout=600)
        if proc.is_alive():
            print("Process stuck, terminating:", proc.pid)
            proc.terminate()
            proc.join()

    successes = 0
    fails = 0
    while not queue.empty():
        _, status = queue.get()
        if status == "success":
            successes += 1
        else:
            fails += 1

    total_seconds = time.time() - start_time

    print("\n==============================")
    print("DONE")
    print("==============================")
    print(f"Success: {successes}")
    print(f"Fail: {fails}")
    print(f"Total time: {total_seconds/60:.2f} minutes")


# ==========================================
# Entry Point
# ==========================================
if __name__ == "__main__":
    import multiprocessing
    multiprocessing.freeze_support()
    main()